import gradio as gr
import pandas as pd
import ollama
import pdfplumber
import easyocr
import cv2
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from PIL import Image

# Initialize OCR with better settings
reader = easyocr.Reader(['hi', 'en'], gpu=False)

def advanced_image_preprocessing(image_path):
    """Enhanced image preprocessing for better OCR accuracy"""
    img = cv2.imread(image_path)
    
    # Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
    # Noise removal using morphological operations
    kernel = np.ones((1, 1), np.uint8)
    gray = cv2.dilate(gray, kernel, iterations=1)
    gray = cv2.erode(gray, kernel, iterations=1)
    
    # Apply Gaussian Blur to reduce noise
    gray = cv2.GaussianBlur(gray, (5, 5), 0)
    
    # Adaptive thresholding for better text extraction
    gray = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
        cv2.THRESH_BINARY, 31, 2
    )
    
    # Deskew image if needed
    coords = np.column_stack(np.where(gray > 0))
    angle = cv2.minAreaRect(coords)[-1]
    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle
    
    if abs(angle) > 0.5:  # Only rotate if needed
        (h, w) = gray.shape[:2]
        center = (w // 2, h // 2)
        M = cv2.getRotationMatrix2D(center, angle, 1.0)
        gray = cv2.warpAffine(gray, M, (w, h), 
                             flags=cv2.INTER_CUBIC, 
                             borderMode=cv2.BORDER_REPLICATE)
    
    return gray

def extract_fir_details_with_patterns(text):
    """Extract FIR details using regex patterns (fallback if AI fails)"""
    details = {
        'fir_number': 'Not Found',
        'section': 'Not Found',
        'date': 'Not Found',
        'accused': 'Not Found',
        'police_station': 'Not Found'
    }
    
    # Pattern for FIR Number
    fir_patterns = [
        r'(?:अपराध|FIR|F\.I\.R|प्र\.सं\.?)\s*(?:क्रमांक|संख्या|नं\.?|No\.?)\s*[:–-]?\s*(\d+[/\-]\d+)',
        r'(?:अपराध|FIR)\s*[:–-]?\s*(\d+[/\-]\d+)',
        r'क्रमांक\s*[:–-]?\s*(\d+[/\-]\d+)'
    ]
    
    # Pattern for Section/धारा
    section_patterns = [
        r'धारा\s*[:–-]?\s*([\d,\s/()]+(?:IPC|आई\.पी\.सी|BNS)?)',
        r'Section\s*[:–-]?\s*([\d,\s/()]+)',
        r'u/s\s*([\d,\s/()]+)'
    ]
    
    # Pattern for Date
    date_patterns = [
        r'दिनांक\s*[:–-]?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
        r'Date\s*[:–-]?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
        r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})'
    ]
    
    # Pattern for Police Station
    ps_patterns = [
        r'थाना\s*[:–-]?\s*([^\n,]+)',
        r'Police\s*Station\s*[:–-]?\s*([^\n,]+)',
        r'P\.S\.?\s*[:–-]?\s*([^\n,]+)'
    ]
    
    # Extract using patterns
    for pattern in fir_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            details['fir_number'] = match.group(1).strip()
            break
    
    for pattern in section_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            details['section'] = match.group(1).strip()
            break
    
    for pattern in date_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            details['date'] = match.group(1).strip()
            break
            
    for pattern in ps_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            details['police_station'] = match.group(1).strip()[:50]
            break
    
    return details

def improve_and_ocr(image_path):
    """Enhanced OCR with better preprocessing and extraction"""
    if image_path is None: 
        return "❌ कोई चित्र अपलोड नहीं किया गया।", "N/A", "N/A", "N/A", "N/A", "N/A"
    
    try:
        # Enhanced preprocessing
        processed_img = advanced_image_preprocessing(image_path)
        
        # Extract text with better parameters
        results = reader.readtext(
            processed_img, 
            detail=0,
            paragraph=True,
            width_ths=0.7,
            height_ths=0.7
        )
        
        full_text = "\n".join(results)
        
        if not full_text.strip():
            return "❌ चित्र से कोई पाठ नहीं निकाला जा सका। कृपया स्पष्ट चित्र अपलोड करें।", "N/A", "N/A", "N/A", "N/A", "N/A"
        
        # First try pattern-based extraction
        pattern_details = extract_fir_details_with_patterns(full_text)
        
        # Then enhance with AI
        prompt = f"""निम्नलिखित FIR पाठ से केवल ये विवरण निकालें। यदि कोई जानकारी नहीं मिलती है तो "नहीं मिला" लिखें।

पाठ: {full_text}

कृपया इन बिंदुओं को निकालें:
1. FIR क्रमांक
2. धारा/Section
3. दिनांक
4. आरोपी का नाम
5. थाना का नाम

केवल संक्षिप्त जानकारी दें, लंबा विवरण न दें।"""
        
        try:
            response = ollama.generate(
                model='llama3.2', 
                prompt=prompt,
                options={
                    'temperature': 0.1,  # More deterministic
                    'top_p': 0.9,
                }
            )
            ai_summary = response['response']
        except Exception as e:
            ai_summary = f"AI विश्लेषण उपलब्ध नहीं है। Pattern-based extraction का उपयोग किया गया।\n\nError: {str(e)}"
        
        # Return structured output
        return (
            full_text[:2000],  # Limit raw text display
            pattern_details['fir_number'],
            pattern_details['section'],
            pattern_details['date'],
            pattern_details['police_station'],
            ai_summary
        )
        
    except Exception as e:
        return f"❌ त्रुटि: {str(e)}", "N/A", "N/A", "N/A", "N/A", "N/A"

def generate_vidhan_sabha(excel_file, pdf_file):
    """Generate Vidhan Sabha report with enhanced formatting"""
    if excel_file is None or pdf_file is None:
        return "❌ कृपया Excel और PDF दोनों फ़ाइलें अपलोड करें।", None
    
    try:
        os.makedirs("data/outputs", exist_ok=True)
        df = pd.read_excel(excel_file.name).fillna("-")
        
        # Extract questions from PDF
        with pdfplumber.open(pdf_file.name) as pdf:
            q_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
        
        if not q_text.strip():
            return "❌ PDF से पाठ निकालने में विफल। कृपया वैध PDF अपलोड करें।", None

        # Create professionally formatted Word document
        doc = Document()
        
        # Add header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "छत्तीसगढ़ पुलिस विभाग"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.runs[0].font.size = Pt(14)
        header_para.runs[0].font.bold = True
        
        # Add title
        title = doc.add_heading('विधान सभा - प्रारूप उत्तर', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generate AI summary
        prompt = f"""निम्नलिखित विधान सभा प्रश्नों के लिए एक औपचारिक हिंदी पुलिस रिपोर्ट सारांश लिखें।

प्रश्न: {q_text[:800]}

डेटा सारांश: {df.head(10).to_string()}

कृपया एक संक्षिप्त, औपचारिक और तथ्यात्मक उत्तर दें (200-300 शब्दों में)।"""
        
        try:
            res = ollama.generate(
                model='llama3.2', 
                prompt=prompt,
                options={'temperature': 0.3}
            )
            summary_text = res['response']
        except Exception as e:
            summary_text = f"डेटा विश्लेषण:\n\nकुल रिकॉर्ड: {len(df)}\n\n[AI सारांश उपलब्ध नहीं: {str(e)}]"
        
        # Add summary
        doc.add_paragraph(summary_text)
        doc.add_paragraph()  # Spacing
        
        # Add data table
        doc.add_heading('परिशिष्ट - आंकड़े तालिका', level=2)
        
        # Create table with styling
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            header_cells[i].text = str(col)
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        
        # Data rows (limit to prevent huge files)
        for idx, row in df.head(50).iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)[:100]  # Limit cell text
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        
        # Add footer note if data was truncated
        if len(df) > 50:
            doc.add_paragraph(f"\n[नोट: तालिका में केवल प्रथम 50 रिकॉर्ड दिखाए गए हैं। कुल रिकॉर्ड: {len(df)}]")
        
        # Save document
        path = "data/outputs/Vidhan_Sabha_Reply.docx"
        doc.save(path)
        
        return summary_text, path
        
    except Exception as e:
        return f"❌ त्रुटि: {str(e)}", None

# --- Professional UI ---
with gr.Blocks(theme=gr.themes.Soft(), title="Police AI System") as demo:
    gr.Markdown("""
    # 🚔 छत्तीसगढ़ पुलिस - AI सहायक प्रणाली
    ### Offline FIR Processing & Report Generation System
    *सभी डेटा स्थानीय रूप से संसाधित - कोई बाहरी API नहीं*
    """)
    
    with gr.Tabs():
        # TAB 1: FIR OCR & Extraction
        with gr.TabItem("📄 FIR OCR & विश्लेषण"):
            gr.Markdown("""
            ### FIR चित्र से स्वचालित विवरण निष्कर्षण
            - उच्च गुणवत्ता वाला FIR चित्र अपलोड करें (JPG/PNG)
            - सिस्टम स्वचालित रूप से महत्वपूर्ण विवरण निकालेगा
            """)
            
            with gr.Row():
                with gr.Column(scale=1):
                    img_in = gr.Image(
                        type="filepath", 
                        label="FIR चित्र अपलोड करें",
                        height=400
                    )
                    btn_ocr = gr.Button("🔍 FIR संसाधित करें", variant="primary", size="lg")
                
                with gr.Column(scale=1):
                    gr.Markdown("#### निकाला गया विवरण:")
                    fir_num = gr.Textbox(label="📋 FIR क्रमांक", interactive=False)
                    section = gr.Textbox(label="⚖️ धारा/Section", interactive=False)
                    date = gr.Textbox(label="📅 दिनांक", interactive=False)
                    ps = gr.Textbox(label="🏢 थाना", interactive=False)
            
            with gr.Accordion("🔎 विस्तृत जानकारी देखें", open=False):
                raw_ocr = gr.TextArea(label="संपूर्ण OCR पाठ", lines=8)
                ai_summary = gr.TextArea(label="AI विश्लेषण", lines=8)
            
            btn_ocr.click(
                improve_and_ocr, 
                inputs=img_in, 
                outputs=[raw_ocr, fir_num, section, date, ps, ai_summary]
            )

        # TAB 2: Vidhan Sabha Reports
        with gr.TabItem("📊 विधान सभा रिपोर्ट"):
            gr.Markdown("""
            ### विधान सभा प्रश्नों के लिए स्वचालित रिपोर्ट
            - अपराध डेटा Excel फ़ाइल अपलोड करें
            - विधान सभा प्रश्न PDF अपलोड करें
            """)
            
            with gr.Row():
                ex_in = gr.File(label="📊 Crime Data (Excel)", file_types=[".xlsx", ".xls"])
                pdf_in = gr.File(label="📄 Questions (PDF)", file_types=[".pdf"])
            
            btn_generate = gr.Button("📝 औपचारिक दस्तावेज़ बनाएं", variant="primary", size="lg")
            
            v_out_txt = gr.TextArea(label="प्रारूप सारांश (हिंदी)", lines=10)
            v_out_file = gr.File(label="📥 रिपोर्ट डाउनलोड करें")
            
            btn_generate.click(
                generate_vidhan_sabha, 
                [ex_in, pdf_in], 
                [v_out_txt, v_out_file]
            )
    
    gr.Markdown("""
    ---
    ### 📌 महत्वपूर्ण सूचना:
    - यह प्रणाली पूर्णतः ऑफलाइन काम करती है
    - कोई डेटा बाहरी सर्वर पर नहीं भेजा जाता
    - Ollama Llama 3.2 मॉडल स्थानीय रूप से चलता है
    - सर्वोत्तम परिणामों के लिए स्पष्ट, उच्च रिज़ॉल्यूशन चित्र अपलोड करें
    """)

if __name__ == "__main__":
    print("🚔 पुलिस AI सिस्टम शुरू हो रहा है...")
    print("📍 निम्न URL पर जाएं:")
    print("   Local: http://localhost:7860/")
    print("   या: http://127.0.0.1:7860/")
    
    demo.launch(
        server_name="127.0.0.1",  # Change to "0.0.0.0" for network access
        server_port=7860,
        show_error=True,
        share=False,
        inbrowser=True
    )